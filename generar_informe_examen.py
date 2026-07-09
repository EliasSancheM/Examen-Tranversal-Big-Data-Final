# -*- coding: utf-8 -*-
"""Genera INFORME_EXAMEN_FINAL.docx — Examen Final Transversal BIY7131.
Fusiona: Prediccion_Transporte_Publico (diseño) + Informe EP2 (batch) +
INFORME_EP3 (streaming), con espacios para las capturas de esta evaluación."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AZUL = RGBColor(0x1F, 0x6F, 0xB2)
DARK = RGBColor(0x20, 0x20, 0x20)
GREY = RGBColor(0x60, 0x60, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

def set_bg(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexc); tcPr.append(shd)

def h(text, size=14, color=AZUL, before=14, after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def para(t, size=11, justify=True, italic=False):
    p = doc.add_paragraph()
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t); r.font.size = Pt(size); r.italic = italic
    return p

def bullet(t):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(t).font.size = Pt(11)

def code(t):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    for ln in t.split("\n"):
        r = p.add_run(ln + "\n"); r.font.name = "Consolas"; r.font.size = Pt(8.5)
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F0F4F8"); pPr.append(shd)

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, x in enumerate(headers):
        set_bg(t.rows[0].cells[i], "1F6FB2")
        run = t.rows[0].cells[i].paragraphs[0].add_run(x)
        run.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = WHITE
    for row in rows:
        c = t.add_row().cells
        for i, v in enumerate(row):
            c[i].paragraphs[0].add_run(v).font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def figura(titulo, captura):
    """Placeholder de captura + pie de figura."""
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"[  INSERTAR CAPTURA:  {captura}  ]")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "FFF7E6"); pPr.append(shd)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(titulo); rc.italic = True; rc.font.size = Pt(8.5); rc.font.color.rgb = GREY

# ============================== PORTADA ==============================
for _ in range(2): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DUOC UC — Escuela de Informática y Telecomunicaciones")
r.bold = True; r.font.size = Pt(12); r.font.color.rgb = GREY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Examen Final Transversal"); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = AZUL
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Plataforma de Datos Batch + Streaming en Google Cloud Platform")
r.font.size = Pt(14); r.font.color.rgb = DARK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RED Santiago — Análisis de Congestión del Transporte Público en Tiempo Real")
r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY
for _ in range(2): doc.add_paragraph()
table(["Campo", "Detalle"], [
    ["Asignatura", "BIY7131 — Big Data"],
    ["Sección", "006D"],
    ["Integrantes", "Benjamín Pumarino · Elías Sánchez · Felipe Villanueva"],
    ["RUT", "[RUT Benjamín] / [RUT Elías] / [RUT Felipe]"],
    ["Docente", "Matías Rojas"],
    ["Ponderación", "40% — Encargo 20% + Presentación 80%"],
    ["Fecha de entrega", "Julio de 2026"],
])
doc.add_page_break()

# ============================== ÍNDICE ==============================
h("Índice")
for i, t in enumerate([
    "Introducción",
    "Justificación de la solución Big Data",
    "Arquitectura, gobierno de datos y ciclo de vida del dato",
    "Procesos, flujos y orquestación (batch + streaming en paralelo)",
    "Conexión con las fuentes de datos",
    "Generación de archivos en el Data Lake",
    "Limpieza, transformación y carga al modelo final",
    "Reportes y panel de control (dashboard)",
    "Conclusiones",
    "Anexos (controles, evidencias y código)",
], 1):
    doc.add_paragraph().add_run(f"{i}.  {t}").font.size = Pt(11)
doc.add_page_break()

# ============================== 1. INTRODUCCIÓN ==============================
h("1. Introducción")
para("El presente informe documenta el proyecto final de la asignatura BIY7131 Big Data: la "
     "construcción de una plataforma de datos que integra procesamiento batch y streaming en "
     "tiempo real sobre Google Cloud Platform. El caso de estudio es la RED de Transporte Público "
     "de Santiago, cuyo objetivo de negocio es analizar la congestión de la flota de buses para "
     "responder preguntas sobre disponibilidad del servicio por ruta, horario y condición climática.")
para("El proyecto consolida el trabajo de las etapas previas: el diseño de la arquitectura y la "
     "justificación Big Data (Etapa 1), el pipeline batch histórico sobre BigQuery (Etapa 2) y la "
     "ingesta streaming mediante una API en Cloud Run (Etapa 3). En esta etapa final, ambos "
     "procesos —batch y streaming— se ejecutan de forma simultánea sobre el mismo Data Lake y el "
     "mismo Data Warehouse, sintetizando la información en un panel de control interactivo.")

# ============================== 2. JUSTIFICACIÓN BIG DATA ==============================
h("2. Justificación de la solución Big Data")
para("La gestión de la RED Santiago cumple con las dimensiones fundamentales del Big Data debido a "
     "la naturaleza masiva y veloz de los datos GPS de la flota. Se justifica el uso de tecnología "
     "cloud y no de una solución tradicional:")
table(["Las 5 V", "Aplicación en el proyecto"], [
    ["Volumen", "Millones de registros GPS generados diariamente por la flota de buses."],
    ["Velocidad", "Las posiciones se actualizan cada pocos segundos; requieren procesamiento inmediato."],
    ["Variedad", "Datos estructurados (CSV GPS), semiestructurados (JSON de clima) y geoespaciales."],
    ["Veracidad", "Manejo de inconsistencias de señal GPS, nulos y registros duplicados."],
    ["Valor", "Transformación de coordenadas crudas en métricas de congestión para decidir refuerzos."],
])
para("Por qué NO otras soluciones:")
bullet("Excel: no escala para procesar millones de registros ni datos en streaming.")
bullet("Cloud SQL (MySQL): servicio de nube, pero no diseñado para analítica masiva ni streaming a gran escala.")
para("Tiempos de respuesta adecuados y manejo de gran variedad de tipos de datos sobre tecnología "
     "cloud (GCP) confirman que Big Data es la solución apropiada.")

# ============================== 3. ARQUITECTURA + GOBIERNO ==============================
h("3. Arquitectura, gobierno de datos y ciclo de vida del dato")
para("Se implementa una Arquitectura Lambda, que combina una capa batch (análisis histórico "
     "profundo) y una capa de velocidad (reacción en tiempo real), unificadas en una capa de "
     "servicio para el dashboard. El mapeo a servicios de Google Cloud es el siguiente:")
table(["Capa Lambda", "Servicio GCP", "Rol en el proyecto"], [
    ["Batch Layer", "Cloud Storage + BigQuery", "Ingesta y modelado del histórico GPS + clima."],
    ["Speed Layer", "Cloud Run (API Flask)", "Ingesta streaming de posiciones GPS en vivo."],
    ["Serving Layer", "BigQuery + Looker Studio", "Vistas analíticas y panel de control interactivo."],
])
para("Análisis de la arquitectura (aspectos característicos y sus beneficios):")
bullet("Escalabilidad: Cloud Run auto-escala con la demanda; BigQuery procesa grandes volúmenes sin administrar servidores.")
bullet("Calidad de datos: validación de rangos y deduplicación en cada capa antes de almacenar.")
bullet("Optimización del procesamiento: transformaciones en SQL sobre BigQuery (procesamiento distribuido).")
bullet("Seguridad y privacidad: autenticación por cuenta de servicio (sin contraseñas); anonimización de identificadores de bus.")
para("Gobierno de datos — roles definidos:")
bullet("Data Owner: define políticas de uso, acceso y seguridad de la información.")
bullet("Data Steward: aplica reglas de validación y controla inconsistencias para asegurar datos confiables.")
bullet("Data Engineer: diseña e implementa los pipelines, el almacenamiento y las transformaciones.")
para("Prácticas de gobierno aplicadas: calidad de datos (filtros de coordenadas fuera del rango de "
     "Santiago), privacidad (anonimización de IDs) y metadata (catálogo de rutas y esquemas). El "
     "ciclo de vida del dato recorre cuatro fases: creación/captura (GPS y clima) → almacenamiento "
     "(Data Lake raw) → procesamiento (limpio y agregado) → consumo/archivado (dashboard y curated).")

# ============================== 4. PROCESOS Y ORQUESTACIÓN ==============================
h("4. Procesos, flujos y orquestación (batch + streaming en paralelo)")
para("El flujo de datos de extremo a extremo integra ambos procesos, ejecutándose de forma "
     "simultánea mediante un orquestador que lanza los dos hilos concurrentes:")
code(
"""Buses GPS ── BATCH  ─▶ Cloud Storage (CSV+JSON) ─▶ BigQuery (raw) ─▶ limpieza/agregación
          └─ STREAM ─▶ API Cloud Run (/api/posicion) ─▶ GCS NDJSON + BigQuery (posiciones_stream)
                                    │
                                    ▼
                     Dataset red_santiago  ─▶  Looker Studio (dashboard)""")
para("El orquestador (orquestador.py) arranca el proceso batch (gcp_pipeline.py) y el proceso "
     "streaming (generador de la API) en el mismo instante; ambos escriben en el dataset "
     "red_santiago en tablas separadas, evitando duplicidad entre capas.")
figura("Figura 1 — Orquestador ejecutando BATCH y STREAMING en paralelo (hilos entrelazados).",
       "log del orquestador con líneas BATCH | ... y STREAM | ... intercaladas")
figura("Figura 2 — Orquestación completada: ambos procesos escribieron en red_santiago.",
       "mensaje 'ORQUESTACIÓN COMPLETADA. Batch y streaming corrieron en paralelo'")

# ============================== 5. CONEXIÓN FUENTES ==============================
h("5. Conexión con las fuentes de datos")
para("Batch — dos fuentes en formatos distintos, cargadas al Data Lake:")
table(["Archivo", "Formato", "Descripción", "Registros"], [
    ["historial_rutas_DTPM.csv", "CSV", "GPS de buses: ID, ruta, velocidad, coordenadas, timestamp", "10.000"],
    ["clima_santiago.json", "JSON", "Clima de Santiago: fecha, condición, temperatura", "31"],
])
para("Streaming — una API propia desplegada en Cloud Run que recibe los pings GPS en línea vía "
     "HTTP POST. La conexión a GCP no usa contraseñas: emplea Application Default Credentials "
     "(cuenta de servicio del servicio de Cloud Run). Parámetros del despliegue:")
table(["Parámetro", "Valor"], [
    ["Proyecto GCP", "qwiklabs-gcp-04-7730387358a9"],
    ["Región", "europe-west1 (Bélgica)"],
    ["Bucket (Data Lake)", "gs://buses-red-examen-2024"],
    ["Dataset BigQuery", "red_santiago"],
    ["Servicio Cloud Run", "red-gps-api"],
    ["URL de la API", "https://red-gps-api-973349414694.europe-west1.run.app"],
])
figura("Figura 3 — Configuración del pipeline batch (Project ID, región y bucket).",
       "grep de gcp_config.py mostrando PROJECT_ID, REGION=europe-west1, BUCKET_NAME")
figura("Figura 4 — Despliegue de la API en Cloud Run y prueba de ingesta (ruta 502c → 502C).",
       "respuesta JSON del POST /api/posicion con status 'aceptado'")

# ============================== 6. DATA LAKE ==============================
h("6. Generación de archivos en el Data Lake")
para("El Data Lake en Cloud Storage organiza los datos por capa y por proceso:")
code(
"""gs://buses-red-examen-2024/
├── datalake/raw/        (BATCH)  historial_rutas_DTPM.csv, clima_ndjson.json
├── datalake/curated/    (BATCH)  datos_enriquecidos.csv, datos_agregados.csv
└── streaming/posiciones/(STREAM) YYYYMMDD/<viaje_id>.json  (NDJSON en vivo)""")
figura("Figura 5 — Estructura del bucket y carga de archivos al Data Lake.",
       "consola de Cloud Storage con los archivos del proyecto")

# ============================== 7. LIMPIEZA/TRANSFORMACIÓN ==============================
h("7. Limpieza, transformación y carga al modelo final")
para("Batch — sobre BigQuery se aplican los seis tratamientos exigidos, produciendo las tablas "
     "datos_enriquecidos (10.000 filas) y datos_agregados (por ruta y hora):")
table(["Tratamiento", "Aplicación en batch"], [
    ["Validación", "Rangos de velocidad (0–120) y coordenadas dentro de Santiago."],
    ["Limpieza", "Imputación de nulos de temperatura con el promedio histórico."],
    ["Deduplicación", "COUNT(DISTINCT ...) y ROW_NUMBER por hash compuesto de campos."],
    ["Normalización", "UPPER(Ruta) y extracción de dimensiones temporales (fecha, hora, día)."],
    ["Enriquecimiento", "LEFT JOIN con clima; categoría de congestión (Fluido/Moderado/Congestionado)."],
    ["Agregación", "Velocidad promedio/std/min/max y total de buses por ruta y hora."],
])
para("Streaming — la API ejecuta en línea, por cada ping, cuatro tratamientos: normalización "
     "(502c → 502C), validación + limpieza + deduplicación (clave bus_id|minuto) y enriquecimiento "
     "(hora, día, nivel de congestión, clima). La agregación se expone en la vista "
     "analytics_congestion_stream.")
para("Métricas de calidad (registros leídos vs. cargados y nulos):")
table(["Proceso", "Leídos", "Cargados", "Nulos", "Observación"], [
    ["Batch GPS (gps_raw → enriquecido)", "10.000", "10.000", "0", "Sin nulos en columnas críticas."],
    ["Batch Clima", "31", "31", "4", "4 nulos de temperatura imputados con el promedio."],
    ["Streaming (API)", "según flujo", "aceptados", "—", "Rechazos por validación y duplicados registrados."],
])
figura("Figura 6 — Control de errores real: detección y corrección del error de particionado FLOAT64.",
       "traceback del error BadRequest FLOAT64 + comando sed de corrección")
figura("Figura 7 — Pipeline batch completado: 10.000 filas enriquecidas y tablas creadas.",
       "salida 'PIPELINE GCP COMPLETADO EXITOSAMENTE' con las tablas del dataset")

# ============================== 8. DASHBOARD ==============================
h("8. Reportes y panel de control (dashboard)")
para("El panel se construyó en Looker Studio conectado a dos fuentes de BigQuery: "
     "analytics_congestion_stream (streaming en vivo) y datos_agregados (histórico batch), con "
     "auto-refresh para reflejar el flujo en tiempo real. Responde directamente a las preguntas de "
     "negocio sobre congestión por ruta, horario y clima.")
table(["Gráfico", "Pregunta que responde", "Insight"], [
    ["Congestión por ruta (barras apiladas)", "¿Qué rutas se congestionan más?", "502 y 210E concentran más tramos congestionados."],
    ["Distribución de congestión (dona)", "¿Cómo se reparte el tráfico?", "65,1% Fluido · 18,6% Congestionado · 16,3% Moderado."],
    ["Pings en vivo (time series)", "¿El streaming está activo?", "La línea crece con cada ping ingerido."],
    ["KPIs (scorecards)", "Estado general", "Total de pings, buses activos y velocidad promedio."],
])
figura("Figura 8 — Conexión de Looker Studio a las tablas del dataset red_santiago.",
       "selector de BigQuery con proyecto, dataset red_santiago y tablas")
figura("Figura 9 — Dashboard final: congestión por ruta, distribución, KPIs y streaming en vivo.",
       "captura del dashboard completo en Looker Studio")

# ============================== 9. CONCLUSIONES ==============================
h("9. Conclusiones")
para("El proyecto demuestra una plataforma Big Data completa y funcional sobre GCP que integra "
     "procesamiento batch y streaming ejecutándose en paralelo, cumpliendo la arquitectura Lambda "
     "propuesta en la etapa de diseño. Se cubrieron la ingesta desde diversas fuentes y formatos, "
     "la limpieza/transformación con los tratamientos exigidos, los controles obligatorios y la "
     "síntesis en un panel de control interactivo.")
bullet("Batch: 10.000 registros GPS + 31 climáticos modelados y agregados en BigQuery.")
bullet("Streaming: API en Cloud Run ingiriendo posiciones en vivo con validación y deduplicación.")
bullet("Ambos procesos poblaron el dataset red_santiago simultáneamente, verificado por conteos > 0.")
bullet("El dashboard revela que ~65% del tráfico circula fluido, con las rutas 502 y 210E como focos de congestión.")

# ============================== ANEXOS ==============================
doc.add_page_break()
h("Anexo A — Controles obligatorios")
table(["Control", "Batch", "Streaming"], [
    ["a) Control de errores", "try/except por etapa; CREATE OR REPLACE (idempotencia).", "try/except en GCS/BigQuery; la API no se cae si un destino falla."],
    ["b) Control de duplicidad", "SELECT DISTINCT + hash compuesto (0 duplicados).", "Deduplicación por clave bus_id|minuto, política SKIP."],
    ["c) Registro de actividad", "estado_ejecuciones_gcp.json + proceso_log + jobs BigQuery.", "api_log (rechazos, duplicados, start/stop)."],
    ["d) Validación / calidad", "Rangos, nulos, leídos vs cargados.", "Validación en línea + vistas v_dq_duplicados / v_dq_velocidad_anomala / v_dq_fuera_santiago."],
])

h("Anexo B — Evidencias adicionales", size=13)
figura("Figura 10 — Descarga de los archivos del proyecto a Cloud Shell (10/10).",
       "salida de gcloud storage cp con 'Completed files 10/10'")
figura("Figura 11 — Tablas y vistas creadas en BigQuery (setup streaming + log).",
       "bq ls red_santiago con posiciones_stream, api_log, proceso_log y vistas v_dq_*")
figura("Figura 12 — Construcción de la imagen Docker (STATUS: SUCCESS).",
       "salida de gcloud builds submit finalizando en SUCCESS")

h("Anexo C — Estructura del repositorio y comandos clave", size=13)
code(
"""Repositorio del proyecto:
  historial_rutas_DTPM.csv / clima_santiago.json   (fuentes batch)
  config.py / gcp_config.py / gcp_pipeline.py      (pipeline batch)
  api_red.py / Dockerfile / requirements.txt        (API streaming Cloud Run)
  setup_bq_red.sql                                  (tablas y vistas streaming)
  orquestador.py                                    (ejecución batch + streaming en paralelo)

Comandos clave (Cloud Shell):
  bq mk --dataset --location=europe-west1 red_santiago
  cat setup_bq_red.sql | bq query --nouse_legacy_sql
  gcloud builds submit --tag gcr.io/$PROJECT_ID/red-gps-api
  gcloud run deploy red-gps-api --image gcr.io/$PROJECT_ID/red-gps-api \\
    --region europe-west1 --allow-unauthenticated \\
    --set-env-vars "GCP_PROJECT=$PROJECT_ID,BQ_DATASET=red_santiago,GCS_BUCKET=buses-red-examen-2024"
  python orquestador.py "$API_URL"     # batch + streaming al mismo tiempo""")

para("Nota: el código fuente completo (config.py, gcp_config.py, gcp_pipeline.py, api_red.py, "
     "orquestador.py, setup_bq_red.sql, Dockerfile y requirements.txt) se entrega como archivos "
     "adjuntos del proyecto, disponibles en el bucket gs://buses-red-examen-2024 y en el "
     "repositorio del equipo.", size=10, italic=True)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— Fin del informe —"); r.font.size = Pt(9.5); r.font.color.rgb = GREY

doc.save("INFORME_EXAMEN_FINAL.docx")

# ---- Documento SEPARADO con el código completo (no cuenta en las 10 páginas) ----
import os as _os
cdoc = Document()
cn = cdoc.styles["Normal"]; cn.font.name = "Consolas"; cn.font.size = Pt(8.5)
cn.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
t0 = cdoc.add_paragraph(); r0 = t0.add_run("Anexo de Código Fuente — Examen Final BIY7131")
r0.bold = True; r0.font.name = "Calibri"; r0.font.size = Pt(15); r0.font.color.rgb = AZUL
CODIGOS = [
    ("config.py — Parámetros de validación", "config.py"),
    ("gcp_config.py — Configuración GCP", "gcp_config.py"),
    ("gcp_pipeline.py — Pipeline BATCH", "gcp_pipeline.py"),
    ("api_red.py — API STREAMING (Cloud Run)", "api_red.py"),
    ("orquestador.py — BATCH + STREAMING en paralelo", "orquestador.py"),
    ("setup_bq_red.sql — Tablas y vistas BigQuery", "setup_bq_red.sql"),
    ("Dockerfile", "Dockerfile"),
    ("requirements.txt", "requirements.txt"),
]
for titulo, path in CODIGOS:
    hp = cdoc.add_paragraph(); hr = hp.add_run(titulo)
    hr.bold = True; hr.font.name = "Calibri"; hr.font.size = Pt(12); hr.font.color.rgb = DARK
    if _os.path.exists(path):
        with open(path, encoding="utf-8") as f:
            for ln in f.read().rstrip("\n").split("\n"):
                cp = cdoc.add_paragraph(); cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.line_spacing = 1.0
                cp.add_run(ln)
    cdoc.add_page_break()
cdoc.save("CODIGO_FUENTE_EXAMEN.docx")
print("CODIGO_FUENTE_EXAMEN.docx generado OK")
print("INFORME_EXAMEN_FINAL.docx generado OK")
