# -*- coding: utf-8 -*-
"""Genera INFORME_EP3.docx (Word) — Evaluación Parcial 3 RED Santiago."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xD6, 0x00, 0x1C)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ---- Estilo base: Arial 11 ----
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def h(text, size=14, color=RED, before=14, after=4, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Arial"; r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def para(text, size=11, justify=True, italic=False):
    p = doc.add_paragraph()
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size); r.italic = italic
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    for ln in text.split("\n"):
        r = p.add_run(ln + "\n")
        r.font.name = "Consolas"; r.font.size = Pt(8)
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    # sombreado del párrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        set_cell_bg(hdr[i], "D6001C")
        pr = hdr[i].paragraphs[0]; run = pr.add_run(htext)
        run.bold = True; run.font.size = Pt(9.5); run.font.color.rgb = WHITE
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            pr = cells[i].paragraphs[0]; run = pr.add_run(val)
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(text).font.size = Pt(11)

# ============================== PORTADA ==============================
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RED Santiago — Análisis de Congestión del\nTransporte Público en Tiempo Real")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RED
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Gestión de Grandes Volúmenes de Datos · Procesamiento Streaming en Google Cloud Platform")
r.font.size = Pt(12); r.font.color.rgb = GREY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Evaluación Parcial N°3 — AVY1101 Big Data"); r.font.size = Pt(11); r.font.color.rgb = GREY
for _ in range(3): doc.add_paragraph()
for linea in [
    "Sección: [COMPLETAR SECCIÓN]",
    "Integrante 1: [Nombre Apellido] — RUT [00.000.000-0]",
    "Integrante 2: [Nombre Apellido] — RUT [00.000.000-0]",
    "Roles simulados: Data Engineers",
    "Fecha: Junio 2026",
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    campo, _, valor = linea.partition(":")
    rb = p.add_run(campo + ": "); rb.bold = True; rb.font.size = Pt(12)
    rv = p.add_run(valor.strip()); rv.font.size = Pt(12)

doc.add_page_break()

# ============================== ÍNDICE ==============================
h("Índice")
for i, t in enumerate([
    "Introducción", "Arquitectura general de la solución",
    "Paso 1 — Conexión a las fuentes de origen",
    "Paso 2 — Ingesta de datos en línea (APIs)",
    "Paso 3 — Limpieza, transformación y carga (ETL Streaming)",
    "Paso 4 — Reportes y Dashboard interactivo",
    "Checklist de controles obligatorios",
    "Conclusiones y próximos pasos",
    "Anexos — Código fuente completo",
], 1):
    p = doc.add_paragraph(); p.add_run(f"{i}.  {t}").font.size = Pt(11)

doc.add_page_break()

# ============================== 1. INTRODUCCIÓN ==============================
h("1. Introducción")
para("El presente informe documenta la construcción de una arquitectura de Big Data en tiempo "
     "real para el monitoreo de la congestión del transporte público de Santiago (buses de la "
     "Red Metropolitana de Movilidad). El objetivo de negocio es responder preguntas sobre la "
     "disponibilidad y fluidez del servicio: ¿en qué rutas y horarios los buses circulan "
     "congestionados?, ¿cuántos buses están activos por ruta?, ¿cómo afecta el clima a la "
     "velocidad de circulación?")
para("A diferencia de la etapa batch (Parcial 2), que procesaba archivos históricos, esta etapa "
     "incorpora ingesta en línea (streaming / real-time): una API desplegada en la nube recibe "
     "los pings GPS de los buses a medida que ocurren, los limpia y transforma en el instante, y "
     "los almacena en un Data Warehouse desde donde un panel interactivo se actualiza "
     "automáticamente. La solución se implementó sobre Google Cloud Platform usando Cloud Run, "
     "Cloud Storage, BigQuery y Looker Studio.")
table(["Componente", "Tecnología GCP", "Rol en el pipeline"], [
    ["Ingesta en línea", "Cloud Run (API Flask)", "Recibe pings GPS vía HTTP en tiempo real"],
    ["Data Lake", "Cloud Storage", "Almacena los eventos crudos en NDJSON"],
    ["Data Warehouse", "BigQuery", "Tablas y vistas analíticas (streaming insert)"],
    ["Panel de control", "Looker Studio", "Dashboard interactivo con auto-refresh"],
])

# ============================== 2. ARQUITECTURA ==============================
h("2. Arquitectura general de la solución")
para("El flujo de datos de extremo a extremo es el siguiente:")
code_block(
"""  Buses (productor de eventos GPS)
            |   POST /api/posicion  (HTTP, en vivo)
            v
  +-----------------------------------------------+
  |   API de Ingesta - Cloud Run (api_red.py)     |
  |   1) Normaliza  2) Valida+Limpia+Deduplica    |
  |   3) Enriquece                                |
  +---------------+---------------+---------------+
                  |               |
                  v               v
   Cloud Storage (Data Lake)   BigQuery (Warehouse)
   streaming/posiciones/*.json  red_santiago.posiciones_stream
                                        |
                                        v  (4) Agregacion (vista SQL)
                                analytics_congestion_stream
                                        |
                                        v
                                Looker Studio (Dashboard)""")
para("Cada posición que llega es escrita en dos destinos simultáneamente: el Data Lake (archivo "
     "NDJSON crudo, para trazabilidad y reproceso) y el Data Warehouse (fila lista para consulta). "
     "Esto garantiza el ciclo de vida del dato: crudo → limpio → agregado → consumo.")
table(["Parámetro", "Valor del despliegue"], [
    ["Proyecto GCP", "qwiklabs-gcp-04-bef2d4403f53"],
    ["Región", "us-west1 (Oregon)"],
    ["Bucket (Data Lake)", "gs://buses_red_2026"],
    ["Dataset BigQuery", "red_santiago"],
    ["Servicio Cloud Run", "red-gps-api"],
    ["URL pública de la API", "https://red-gps-api-809939043386.us-west1.run.app"],
])

# ============================== 3. PASO 1 ==============================
doc.add_page_break()
h("3. Paso 1 — Conexión a las fuentes de origen")
h("Origen de los datos", size=11, color=GREY, before=8)
para("La fuente de origen es la flota de buses, que actúa como productor de eventos. Cada bus "
     "reporta su posición (identificador, ruta, velocidad y coordenadas GPS). Para la demostración, "
     "el productor se simula mediante un generador que emite pings con las mismas características "
     "que la flota real (rutas 210, 210E, 502, 502C, 104; velocidades de 0 a 70 km/h; coordenadas "
     "dentro del bounding box de Santiago). El reemplazo por la API real de la flota (GTFS-Realtime) "
     "solo requeriría cambiar la función generadora; el resto del pipeline permanece idéntico.")
h("Cadena de conexión y credenciales (sin exponer contraseñas)", size=11, color=GREY, before=8)
para("La conexión a GCP no usa contraseñas: emplea Application Default Credentials, es decir, la "
     "identidad de la cuenta de servicio asignada al servicio de Cloud Run. Los parámetros de "
     "conexión se inyectan como variables de entorno en el despliegue:")
code_block('gcloud run deploy red-gps-api \\\n  --set-env-vars "GCP_PROJECT=$PROJECT_ID,BQ_DATASET=red_santiago,GCS_BUCKET=buses_red_2026,..."')
para("Dentro del código, los clientes se construyen con esos valores (clientes perezosos):")
code_block('def bq():\n    if _bq_client is None and _GCP_OK and PROJECT_ID:\n        _bq_client = bigquery.Client(project=PROJECT_ID)   # autenticacion por ADC\n    return _bq_client')
h("Controles aplicados en este paso", size=11, color=GREY, before=8)
table(["Control", "Implementación"], [
    ["Control de errores", "Si faltan las librerías GCP o el PROJECT_ID, la API arranca igual en modo memoria y no se cae (try/except ImportError)."],
    ["Validación", "Se verifica que exista proyecto y dataset antes de intentar conectar."],
    ["Registro de actividad", "El endpoint raíz / expone proyecto, dataset y bucket conectados para verificar la conexión."],
])

# ============================== 4. PASO 2 ==============================
h("4. Paso 2 — Ingesta de datos en línea (APIs)")
para("La ingesta se realiza mediante una API propia construida con Flask y desplegada en Cloud "
     "Run, que expone endpoints HTTP para recibir los datos en línea. Esto cumple el requisito de "
     "usar APIs de la industria/propias para ingesta en tiempo real.")
table(["Endpoint", "Método", "Función"], [
    ["/api/posicion", "POST", "Ingesta un ping GPS individual"],
    ["/api/auto/start", "POST", "Activa el generador continuo de eventos"],
    ["/api/auto/stop", "POST", "Detiene el generador"],
    ["/api/congestion/resumen", "GET", "Consulta agregada en vivo desde BigQuery"],
    ["/api/metricas", "GET", "Contadores de ingesta (recibidos/aceptados/rechazados)"],
])
para("La imagen del servicio se construye con Cloud Build a partir del Dockerfile y se publica "
     "con gcloud run deploy. Cada ping aceptado se persiste en el Data Lake (Cloud Storage, NDJSON) "
     "y en BigQuery (streaming insert con insert_rows_json).")
h("Evidencia — ingesta de un ping real (respuesta de la API)", size=11, color=GREY, before=8)
code_block(
"""POST /api/posicion  {"bus_id":"BUS_1040","ruta":"502c","velocidad_kmh":12,"latitud":-33.45,"longitud":-70.66}

{"status":"aceptado","registro":{
   "bus_id":"BUS_1040","ruta":"502C","ruta_raw":"502c","velocidad_kmh":12.0,
   "hora":3,"dia_semana":"Saturday","nivel_congestion":"Congestionado",
   "clima_condicion":"Despejado","clima_temperatura":14.0,
   "viaje_id":"cd904fbe-22ce-44fc-8d56-868cf516255f"}}""")
h("Controles aplicados en este paso", size=11, color=GREY, before=8)
table(["Control", "Implementación"], [
    ["Control de errores", "La escritura a GCS y BigQuery va en try/except: si un destino falla, se registra un warning y la API sigue operando sin perder el evento en el otro destino."],
    ["Registro de actividad", "Cada acción se escribe en la tabla de logs red_santiago.api_log (start/stop, rechazos, duplicados)."],
    ["Validación", "El endpoint exige JSON con campos obligatorios; los inválidos retornan HTTP 422."],
])

# ============================== 5. PASO 3 ==============================
doc.add_page_break()
h("5. Paso 3 — Limpieza, transformación y carga (ETL Streaming)")
para("Todo el ETL se ejecuta en línea, evento por evento, antes de almacenar. El proceso "
     "implementa explícitamente los seis tratamientos exigidos por la rúbrica:")
table(["Tratamiento", "Dónde", "Qué hace"], [
    ["Normalización", "normalizar()", "Ruta a forma canónica (502c → 502C), bus_id en mayúsculas, velocidad redondeada."],
    ["Validación", "validar()", "Velocidad en rango 0–120 km/h; coordenadas dentro de Santiago. Fuera de rango ⇒ rechazo 422."],
    ["Limpieza", "validar()", "Descarta registros corruptos o incompletos antes de persistir."],
    ["Deduplicación", "es_duplicado()", "Un mismo bus no puede registrar dos pings en el mismo minuto (clave bus_id|minuto)."],
    ["Enriquecimiento", "enriquecer()", "Agrega hora, dia_semana, nivel_congestion y datos de clima."],
    ["Agregación", "Vista SQL", "analytics_congestion_stream agrupa por ruta, hora y nivel de congestión."],
])
h("Trazabilidad y ciclo de vida del dato", size=11, color=GREY, before=8)
para("Cada registro recibe un identificador único viaje_id (UUID) y un ingest_timestamp, lo que "
     "permite rastrear cualquier fila desde su origen hasta el dashboard. Se conserva además el "
     "campo ruta_raw (valor original recibido) junto a la ruta normalizada, evidenciando la "
     "transformación aplicada. El dato transita por las capas raw (GCS) → posiciones_stream "
     "(limpio) → analytics_congestion_stream (agregado).")
h("Evidencia — control de errores y de duplicidad en vivo", size=11, color=GREY, before=8)
code_block(
"""// Validacion: velocidad fuera de rango -> rechazado
POST {"bus_id":"BUS_999","ruta":"104","velocidad_kmh":999,...}
{"status":"rechazado","motivo":"velocidad fuera de rango (999.0)"}

// Deduplicacion: segundo ping del mismo bus/minuto -> rechazado
POST {"bus_id":"BUS_2020","ruta":"210",...}   -> {"status":"aceptado",...}
POST {"bus_id":"BUS_2020","ruta":"210",...}   -> {"status":"rechazado","motivo":"duplicado (mismo bus, mismo minuto)"}""")
h("Decisión sobre duplicados (política)", size=11, color=GREY, before=8)
para("Ante un dato que ya existe, la política adoptada es SKIP (descartar): el evento duplicado se "
     "rechaza y se registra en el log, en lugar de hacer upsert. Se eligió SKIP porque en "
     "telemetría GPS un segundo reporte idéntico en el mismo minuto es ruido, no una actualización "
     "de estado; conservarlo distorsionaría los promedios de velocidad y los conteos.")

# ============================== 6. PASO 4 ==============================
h("6. Paso 4 — Reportes y Dashboard interactivo")
para("El panel se construyó en Looker Studio conectado a las vistas de BigQuery, con auto-refresh "
     "para reflejar el streaming. Responde directamente a las preguntas de negocio sobre "
     "disponibilidad y congestión del servicio.")
table(["Gráfico", "Pregunta de negocio que responde", "Dimensión / Métrica"], [
    ["Barras apiladas", "¿Qué rutas están más congestionadas?", "ruta × nivel_congestion / total_pings"],
    ["Línea temporal", "¿En qué horarios baja la velocidad?", "hora / velocidad_promedio"],
    ["Barras", "¿Cómo se distribuye la congestión?", "nivel_congestion / total_pings"],
    ["Mapa geográfico", "¿Dónde están los buses ahora?", "latitud, longitud / velocidad"],
    ["Tabla", "¿Cómo impacta el clima en la circulación?", "clima_condicion / velocidad_promedio"],
    ["Scorecards", "KPIs globales en vivo", "total_pings, buses_unicos, vel. prom."],
])
p = doc.add_paragraph()
r = p.add_run("⚠ PENDIENTE: reemplazar los recuadros siguientes por las capturas reales del dashboard antes de entregar.")
r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RED
for txt in ["[ Captura 1 — Vista general del dashboard con scorecards y filtros ]",
            "[ Captura 2 — Barras apiladas de congestión por ruta + línea de velocidad por hora ]"]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(20)

# ============================== 7. CHECKLIST ==============================
doc.add_page_break()
h("7. Checklist de controles obligatorios")
table(["Pregunta de control", "Respuesta de nuestra solución"], [
    ["¿Qué pasa si falla la conexión? (Control de errores)",
     "Cada escritura (GCS / BigQuery) está envuelta en try/except. Si un destino falla, se registra un warning y el evento se conserva en el otro destino; la API nunca se cae. Si faltan credenciales, opera en modo memoria."],
    ["¿Qué pasa si la API trae datos ya guardados? (Control de duplicidad)",
     "Política SKIP: se deduplica por clave bus_id|minuto. El duplicado se rechaza con HTTP 422 y se registra en api_log. La vista v_dq_duplicados permite auditar duplicados residuales."],
    ["¿Dónde se guardan los logs? (Registro de actividad)",
     "En la tabla red_santiago.api_log (event_id, tipo, descripción, estado, timestamp) y en Cloud Logging. El endpoint /api/metricas expone contadores en vivo."],
    ["¿Cómo sabemos que los datos transformados están correctos? (Validación)",
     "Validación de rangos en línea + tres vistas de calidad en BigQuery: v_dq_duplicados, v_dq_velocidad_anomala y v_dq_fuera_santiago, que deben retornar 0 filas."],
])

# ============================== 8. CONCLUSIONES ==============================
h("8. Conclusiones y próximos pasos")
para("La arquitectura implementada demuestra un pipeline de ingesta en línea de extremo a extremo: "
     "los datos se reciben, limpian, transforman y almacenan en tiempo real, y se sintetizan en un "
     "panel que un gerente puede usar para decidir refuerzos de flota por ruta y horario. Se "
     "cubrieron los seis tratamientos de datos exigidos y los cuatro controles obligatorios, con "
     "evidencia funcional verificada en GCP.")
para("Próximos pasos: (1) conectar la API real de la flota (GTFS-Realtime) reemplazando solo el "
     "generador; (2) particionar posiciones_stream por fecha para optimizar costos; (3) agregar "
     "alertas automáticas cuando una ruta supere un umbral de congestión sostenido.")

# ============================== 9. ANEXOS ==============================
doc.add_page_break()
h("9. Anexos — Código fuente completo")
anexos = [
    ("Anexo A — api_red.py (API de ingesta en streaming, Cloud Run)", "api_red.py"),
    ("Anexo B — Dockerfile", "Dockerfile"),
    ("Anexo C — requirements.txt", "requirements.txt"),
    ("Anexo D — setup_bq_red.sql (tablas y vistas en BigQuery)", "setup_bq_red.sql"),
    ("Anexo E — demo_stream_red.sh (generador de tráfico para la demo)", "demo_stream_red.sh"),
]
for titulo, path in anexos:
    h(titulo, size=12, color=RGBColor(0x33,0x33,0x33), before=12)
    with open(path, encoding="utf-8") as f:
        code_block(f.read().rstrip("\n"))

h("Anexo F — Comandos de despliegue en GCP (Cloud Shell)", size=12, color=RGBColor(0x33,0x33,0x33), before=12)
code_block(
"""# 1. Variables
PROJECT_ID="$(gcloud config get-value project)"
REGION="us-west1"
BUCKET_NAME="buses_red_2026"

# 2. Habilitar servicios
gcloud services enable run.googleapis.com cloudbuild.googleapis.com \\
  bigquery.googleapis.com storage.googleapis.com

# 3. Data Lake
gcloud storage buckets create "gs://$BUCKET_NAME" --location="$REGION" \\
  --uniform-bucket-level-access

# 4. Dataset + tablas/vistas streaming
bq mk --dataset --location="$REGION" red_santiago
cat setup_bq_red.sql | bq query --nouse_legacy_sql

# 5. Build + Deploy
gcloud builds submit --tag "gcr.io/$PROJECT_ID/red-gps-api"
gcloud run deploy red-gps-api \\
  --image "gcr.io/$PROJECT_ID/red-gps-api" --platform managed --region "$REGION" \\
  --allow-unauthenticated \\
  --set-env-vars "GCP_PROJECT=$PROJECT_ID,BQ_DATASET=red_santiago,GCS_BUCKET=$BUCKET_NAME,AUTO_INTERVAL=2,BUSES_POR_TICK=20" \\
  --min-instances 1 --timeout 300

# 6. Encender flujo continuo
API_URL=$(gcloud run services describe red-gps-api --region "$REGION" --format="value(status.url)")
curl -s -X POST "$API_URL/api/auto/start" """)

doc.save("INFORME_EP3.docx")
print("INFORME_EP3.docx generado OK")
